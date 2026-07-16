"""
Word 文档编辑器 — 在原始 .docx 上做精确文本替换，保留原始排版
"""

from pathlib import Path
from typing import Optional
from copy import deepcopy

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn


def _get_font_style(run) -> dict:
    """获取 run 的字体样式"""
    font = run.font
    style = {
        "bold": font.bold,
        "italic": font.italic,
        "size": font.size,
        "color": font.color.rgb if font.color and font.color.rgb else None,
        "name": font.name,
        "underline": font.underline,
    }
    return style


def _apply_font_style(new_run, style: dict):
    """将字体样式应用到新的 run"""
    if style["bold"] is not None:
        new_run.bold = style["bold"]
    if style["italic"] is not None:
        new_run.italic = style["italic"]
    if style["size"] is not None:
        new_run.font.size = style["size"]
    if style["color"] is not None:
        new_run.font.color.rgb = style["color"]
    if style["name"] is not None:
        new_run.font.name = style["name"]
    if style["underline"] is not None:
        new_run.underline = style["underline"]


def apply_replacements(
    docx_path: str,
    replacements: list[dict],
    output_path: str,
) -> str:
    """
    在原始 Word 文档中执行精确文本替换，保留所有原始排版。

    Args:
        docx_path: 原始 .docx 文件路径
        replacements: [{"original": "原文", "suggested": "替换文本"}, ...]
        output_path: 输出文件路径

    Returns:
        output_path
    """
    doc = Document(docx_path)

    for replacement in replacements:
        old_text = replacement["original"]
        new_text = replacement["suggested"]
        if not old_text or not new_text or old_text == new_text:
            continue

        # 第一步：在段落中查找替换
        for para in doc.paragraphs:
            _replace_in_paragraph(para, old_text, new_text)

        # 第二步：在表格中查找替换
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _replace_in_paragraph(para, old_text, new_text)

    doc.save(output_path)
    return output_path


def _replace_in_paragraph(para, old_text: str, new_text: str) -> bool:
    """
    在单个段落中执行文本替换，保留原始排版。
    返回是否找到并替换了文本。
    """
    if old_text not in para.text:
        return False

    # 获取所有 run 的文本组合
    full_text = "".join(run.text for run in para.runs)
    if old_text not in full_text:
        return False

    # 找到 old_text 在全文中的位置
    start_idx = full_text.find(old_text)
    if start_idx == -1:
        return False
    end_idx = start_idx + len(old_text)

    # 计算每个 run 的累积位置
    run_positions = []
    pos = 0
    for run in para.runs:
        run_len = len(run.text)
        run_positions.append((run, pos, pos + run_len))
        pos += run_len

    # 找到覆盖范围的 run
    affected_runs = []
    for run, r_start, r_end in run_positions:
        if r_start < end_idx and r_end > start_idx:
            affected_runs.append((run, r_start, r_end))

    if not affected_runs:
        return False

    # 获取第一个 run 的样式（用于替换文本）
    first_run = affected_runs[0][0]
    style = _get_font_style(first_run)

    # 情况1：old_text 完全在一个 run 内
    if len(affected_runs) == 1:
        run, r_start, r_end = affected_runs[0]
        before = run.text[:start_idx - r_start]
        after = run.text[end_idx - r_start:]
        run.text = before + new_text + after
        return True

    # 情况2：跨多个 run
    # 第一个 run：保留 before 部分
    first, fs_start, fs_end = affected_runs[0]
    before_text = first.text[:start_idx - fs_start]
    first.text = before_text + new_text

    # 中间的 run：清空
    for i in range(1, len(affected_runs) - 1):
        affected_runs[i][0].text = ""

    # 最后一个 run：保留 after 部分
    last, ls_start, ls_end = affected_runs[-1]
    after_text = last.text[end_idx - ls_start:]
    last.text = after_text

    return True


def convert_to_pdf(docx_path: str, pdf_path: str) -> str:
    """
    使用 Word COM 将 .docx 转换为 PDF，保留完整排版。
    仅在 Windows 上可用（需要安装 Microsoft Word）。

    Args:
        docx_path: 输入 .docx 路径
        pdf_path: 输出 .pdf 路径

    Returns:
        pdf_path
    """
    import win32com.client
    import pythoncom

    # 初始化 COM
    pythoncom.CoInitialize()

    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False

        abs_docx = str(Path(docx_path).resolve())
        abs_pdf = str(Path(pdf_path).resolve())

        doc = word.Documents.Open(abs_docx)
        # 17 = wdFormatPDF
        doc.SaveAs(abs_pdf, FileFormat=17)
        doc.Close()

        return pdf_path
    finally:
        if doc is not None:
            try:
                doc.Close()
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()